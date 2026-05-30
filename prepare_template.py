# -*- coding: utf-8 -*-
"""
Convierte el reporte mensual original
"Monthly Reconciliation Diesel Report From <Month> 1st to <Day> <Year>.docx"
en una PLANTILLA docxtpl con etiquetas Jinja ({{ ... }} y {%tr ... %}).

Uso:
    python prepare_template.py "ruta/al/reporte_original.docx"

Genera:  plantilla_reporte.docx

Mapeo de etiquetas (basado en el reporte de Abril 2026):
  - Tabla 0: consolidada del LFO (1 fila)
  - Tabla 1: deliveries diarias (loop + Grand Total)
  - Tabla 2: service trucks (loop)
  - Tabla 3: weekly breakdown del mes (loop)
  - Tabla 4: Findings — SE DEJA INTACTA. La llena un tercero a mano.

NOTA: las posiciones de parrafo y tabla estan calibradas para el reporte de
Abril 2026 (223 parrafos, 5 tablas). Marzo 2026 tiene la misma estructura
con N filas distinta en Tabla 4 (que no se modifica), por lo que el script
funciona igual.
"""
from __future__ import annotations

import copy
import os
import re
import sys
import zipfile

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

# Namespace del DrawingML Chart (graficos nativos de Word).
NS_CHART = "http://schemas.openxmlformats.org/drawingml/2006/chart"

DEFAULT_FONT = "Arial"
OUTPUT_NAME = "plantilla_reporte.docx"

# Detecta frases de periodo y fechas sueltas en parrafos.
# El reporte mensual usa formatos como "1st April to 30th 2026" o
# "From 1st April to 30th 2026".
PERIOD_RE = re.compile(
    r"\b\d{1,2}(?:st|nd|rd|th)?\s+[A-Z][a-z]+\s+to\s+\d{1,2}(?:st|nd|rd|th)?"
    r"\s+\d{4}\b"
)
PERIOD_RE_WEEKLY = re.compile(
    r"[A-Z][a-z]+\s+\d{1,2}(?:st|nd|rd|th)?\s+to\s+"
    r"(?:[A-Z][a-z]+\s+)?\d{1,2}(?:st|nd|rd|th)?\s*,?\s*\d{4}")
STANDALONE_DATE_RE = re.compile(
    r"[A-Z][a-z]+\s+\d{1,2}(?:st|nd|rd|th)?\s*,?\s*\d{4}")
# Para captions tipo "Inlet vs Delivery tickets for April, 2026."
MONTH_YEAR_RE = re.compile(
    r"\b[A-Z][a-z]+,?\s*\d{4}\b")

# Indices del reporte original entregado el 20260501 (Abril 2026).
# Las imagenes con drawing en parrafo:
#  0,6,17 = covers/decorativas (no tocar)
#  61     = icono pequeno (no tocar)
#  67     = Figura 1 (deliveries daily bars)
#  92     = Figura 2 (donut)
#  107    = Figura 3 (daily tank log)
#  113    = Figura 4 (monthly recon trend)
#  125    = Figura 5 (delivery vs recon)
#  130    = Figura 6 (monthly breakdown bars)
#  133    = Figura 7 (historical outflows)
#  179    = fig_tasks
FIGURE_PARAGRAPHS = {
    67: "fig1",
    92: "fig2",
    107: "fig3",
    113: "fig4",
    125: "fig5",
    130: "fig6",
    133: "fig7",
    179: "fig_tasks",
}

# Parrafos de texto que se reemplazan completos por una etiqueta docxtpl.
TEXT_PARAGRAPHS = {
    42: "From {{ period_full }}",
    45: "{{r recon_sentence }}",
    58: "{{r delivery_narrative }}",
    60: "{{r delivery_band_line }}",
    65: "{{r delivery_stats_line }}",
    99: "{{r disp_total_line }}",
    101: "{{r disp_equipment_line }}",
    102: "{{r disp_other_line }}",
    103: "{{r disp_transfers_line }}",
    120: "{{r monthly_transactions_line }}",
    121: "{{r monthly_to_equipment_line }}",
    122: "{{r monthly_other_line }}",
    123: "{{r monthly_transfers_line }}",
}

# Grupo de "Considerations" (49-52) -> bucle Jinja.
CONSIDERATION_GROUPS = {
    "considerations": [49, 50, 51, 52],
}

# Tablas del reporte (orden por aparicion en el DOCX):
#   0 -> Consolidada LFO
#   1 -> Deliveries daily (Table general)
#   2 -> Service Trucks
#   3 -> Weekly Breakdown (Week, Date, Transactions, To Equipment, Transfers)
#   4 -> Findings (NO se toca)


def _rewrite_keep_format(paragraph, new_text):
    runs = paragraph.runs
    bold = italic = underline = None
    name = None
    size = None
    if runs:
        r0 = runs[0]
        bold, italic, underline = r0.bold, r0.italic, r0.underline
        name = r0.font.name
        size = r0.font.size
    for child in list(paragraph._p):
        if child.tag in (qn("w:r"), qn("w:hyperlink")):
            paragraph._p.remove(child)
    run = paragraph.add_run(new_text)
    run.bold, run.italic, run.underline = bold, italic, underline
    run.font.name = name or DEFAULT_FONT
    if size is not None:
        run.font.size = size


def set_paragraph_text(paragraph, text):
    _rewrite_keep_format(paragraph, text)


def delete_paragraph(paragraph):
    el = paragraph._p
    el.getparent().remove(el)


def retag_period(paragraph):
    """Sustituye fechas de periodo por {{ period_full }} en captions."""
    text = paragraph.text
    if not text or "{{" in text or "{%" in text:
        return False
    # Prioridad: PERIOD_RE mensual ("1st April to 30th 2026").
    m = PERIOD_RE.search(text)
    if m:
        _rewrite_keep_format(paragraph,
                              PERIOD_RE.sub("{{ period_full }}", text))
        return True
    return False


def retag_month_year(paragraph):
    """En captions de figuras del tipo 'for April, 2026.' o
    '- April 2026.' sustituye por '{{ month_label }}'."""
    text = paragraph.text
    if not text or "{{" in text or "{%" in text:
        return False
    m = MONTH_YEAR_RE.search(text)
    if not m:
        return False
    # Evitar sustituir 'Monthly' o palabras de un solo nombre por error.
    matched = m.group(0)
    # Solo aplicar si tiene un ano (4 digitos).
    if not re.search(r"\d{4}", matched):
        return False
    _rewrite_keep_format(paragraph,
                          MONTH_YEAR_RE.sub("{{ month_label }}", text))
    return True


def retag_textbox_dates(doc):
    """Etiqueta fechas en cuadros de texto (cubierta/portada)."""
    count = 0
    for txbx in doc.element.iter(qn("w:txbxContent")):
        for p_el in txbx.iter(qn("w:p")):
            para = Paragraph(p_el, doc)
            text = para.text
            if not text.strip() or "{{" in text or "{%" in text:
                continue
            # Caso 1: contiene un periodo.
            m1 = PERIOD_RE.search(text) or PERIOD_RE_WEEKLY.search(text)
            if m1:
                _rewrite_keep_format(
                    para, m1.re.sub("{{ period_full }}", text))
                count += 1
                continue
            # Caso 2: fecha suelta "April 30, 2026" -> {{ cover_date }}.
            if STANDALONE_DATE_RE.search(text):
                _rewrite_keep_format(
                    para, STANDALONE_DATE_RE.sub(
                        "{{ cover_date }}", text))
                count += 1
                continue
            # Caso 3: "April 2026" -> {{ month_label }}.
            if MONTH_YEAR_RE.search(text):
                _rewrite_keep_format(
                    para, MONTH_YEAR_RE.sub(
                        "{{ month_label }}", text))
                count += 1
    return count


def set_cell(cell, text):
    cell.text = text


def _insert_loop_rows(table, data_row, for_text, end_text):
    for_tr = copy.deepcopy(data_row._tr)
    end_tr = copy.deepcopy(data_row._tr)
    data_row._tr.addprevious(for_tr)
    data_row._tr.addnext(end_tr)
    for tr, text in ((for_tr, for_text), (end_tr, end_text)):
        cells = tr.findall(".//" + qn("w:tc"))
        for ci, tc in enumerate(cells):
            for para in tc.findall(qn("w:p")):
                for child in list(para):
                    if child.tag in (qn("w:r"), qn("w:hyperlink")):
                        para.remove(child)
            if ci == 0:
                para = tc.findall(qn("w:p"))[-1]
                run = para.makeelement(qn("w:r"), {})
                wt = para.makeelement(qn("w:t"), {})
                wt.text = text
                run.append(wt)
                para.append(run)


def tag_consolidated_table(table):
    """Tabla 0: Site | Opening | Deliveries | Transactions | Calculated
    | Net Stock Change | Closing | Variance | %  (1 fila de datos)."""
    fields = ["site", "opening", "deliveries", "transactions",
              "calc_stock", "net_change", "closing", "variance", "pct"]
    if len(table.rows) < 2:
        return
    row = table.rows[1]
    for i, field in enumerate(fields):
        if i < len(row.cells):
            set_cell(row.cells[i], "{{ cons.%s }}" % field)


def tag_delivery_table(table):
    """Tabla 1: '' | Delivery Tickets | Inlet Deliveries | Sum of Variance
    | Variance %.

    Loop de filas + fila final Grand Total."""
    fields = ["date", "tickets", "inlet", "variance", "pct"]
    if len(table.rows) < 3:
        return
    data_row = table.rows[1]
    _insert_loop_rows(table, data_row,
                      "{%tr for d in deliveries %}", "{%tr endfor %}")
    for i, field in enumerate(fields):
        if i < len(data_row.cells):
            set_cell(data_row.cells[i], "{{ d.%s }}" % field)
    # Eliminar filas intermedias (las del template original con datos).
    # Tras _insert_loop_rows el orden es:
    #   [header, for, data_row, endfor, daily_2, daily_3, ..., Grand Total]
    # Conservamos header (R0), for (R1), data_row (R2), endfor (R3) y la
    # ultima fila (Grand Total).
    while len(table.rows) > 5:
        # Elimina la fila en indice 4 (despues de endfor) hasta dejar solo
        # la ultima (Grand Total).
        table._tbl.remove(table.rows[4]._tr)
    total_row = table.rows[-1]
    if len(total_row.cells) >= 5:
        set_cell(total_row.cells[0], "Total general")
        set_cell(total_row.cells[1], "{{ dtot.tickets }}")
        set_cell(total_row.cells[2], "{{ dtot.inlet }}")
        set_cell(total_row.cells[3], "{{ dtot.variance }}")
        set_cell(total_row.cells[4], "{{ dtot.pct }}")


def tag_trucks_table(table):
    """Tabla 2: Site | Deliveries | Transactions."""
    fields = ["site", "deliveries", "transactions"]
    if len(table.rows) < 2:
        return
    data_row = table.rows[1]
    _insert_loop_rows(table, data_row,
                      "{%tr for t in trucks %}", "{%tr endfor %}")
    for i, field in enumerate(fields):
        if i < len(data_row.cells):
            set_cell(data_row.cells[i], "{{ t.%s }}" % field)
    # Tras _insert_loop_rows el orden es:
    #   [header, for, data_row, endfor, truck2, truck3, ...]
    for extra in list(table.rows[4:]):
        table._tbl.remove(extra._tr)


def tag_weekly_table(table):
    """Tabla 3: Week | Date | Transactions | To Equipment | Transfers."""
    fields = ["week", "date", "transactions", "to_equipment", "transfers"]
    if len(table.rows) < 2:
        return
    data_row = table.rows[1]
    _insert_loop_rows(table, data_row,
                      "{%tr for w in weekly %}", "{%tr endfor %}")
    for i, field in enumerate(fields):
        if i < len(data_row.cells):
            set_cell(data_row.cells[i], "{{ w.%s }}" % field)
    for extra in list(table.rows[4:]):
        table._tbl.remove(extra._tr)


def strip_native_charts(doc) -> int:
    """Elimina del cuerpo del documento cualquier run que contenga un grafico
    NATIVO de Word (<c:chart>), junto con su <w:drawing>.

    El reporte manual trae graficos nativos con hojas de Excel embebidas
    (word/charts/* + word/embeddings/*). Esos objetos, al pasar por
    python-docx, son la causa tipica del error de Word 'unreadable content'.
    Como el software genera TODAS las figuras con matplotlib (imagenes PNG),
    los graficos nativos sobran y se eliminan."""
    body = doc.element.body
    removed = 0
    for drawing in list(body.iter(qn("w:drawing"))):
        if drawing.find(".//{%s}chart" % NS_CHART) is None:
            continue
        # Subir hasta el <w:r> contenedor y removerlo.
        run = drawing
        while run is not None and run.tag != qn("w:r"):
            run = run.getparent()
        target = run if run is not None else drawing
        parent = target.getparent()
        if parent is not None:
            parent.remove(target)
            removed += 1
    return removed


def purge_chart_parts(docx_path: str) -> dict:
    """Cirugia sobre el .docx (zip) para eliminar TODO rastro de graficos
    nativos y hojas Excel embebidas:
      - parts bajo word/charts/ y word/embeddings/ (y sus _rels).
      - relaciones en word/_rels/document.xml.rels que apunten a charts.
      - Overrides de [Content_Types].xml de las partes de chart.
    Idempotente: si no hay charts, no hace nada.

    Devuelve un resumen {dropped_parts, dropped_rels, dropped_overrides}."""
    with zipfile.ZipFile(docx_path, "r") as zin:
        names = zin.namelist()
        contents = {n: zin.read(n) for n in names}

    drop_parts = set()
    for n in names:
        low = n.lower()
        if low.startswith("word/charts/") or low.startswith("word/embeddings/"):
            drop_parts.add(n)

    if not drop_parts:
        return {"dropped_parts": 0, "dropped_rels": 0, "dropped_overrides": 0}

    # 1) Limpiar document.xml.rels: quitar relaciones a charts/embeddings.
    rels_name = "word/_rels/document.xml.rels"
    dropped_rels = 0
    if rels_name in contents:
        text = contents[rels_name].decode("utf-8")
        # Cada <Relationship .../> en una sola etiqueta.
        def _keep(match):
            nonlocal dropped_rels
            tag = match.group(0)
            tgt_m = re.search(r'Target="([^"]+)"', tag)
            tgt = (tgt_m.group(1) if tgt_m else "").lower()
            if "charts/" in tgt or "embeddings/" in tgt:
                dropped_rels += 1
                return ""
            return tag
        text = re.sub(r"<Relationship\b[^>]*?/>", _keep, text)
        contents[rels_name] = text.encode("utf-8")

    # 2) Limpiar [Content_Types].xml: quitar Overrides de chart parts.
    ct_name = "[Content_Types].xml"
    dropped_overrides = 0
    if ct_name in contents:
        text = contents[ct_name].decode("utf-8")
        def _keep_ct(match):
            nonlocal dropped_overrides
            tag = match.group(0)
            pn_m = re.search(r'PartName="([^"]+)"', tag)
            pn = (pn_m.group(1) if pn_m else "").lower()
            if pn.startswith("/word/charts/") or \
                    pn.startswith("/word/embeddings/"):
                dropped_overrides += 1
                return ""
            return tag
        text = re.sub(r"<Override\b[^>]*?/>", _keep_ct, text)
        contents[ct_name] = text.encode("utf-8")

    # 3) Reescribir el zip sin las partes de chart/embeddings ni sus _rels.
    def _is_dropped(name: str) -> bool:
        if name in drop_parts:
            return True
        low = name.lower()
        # _rels de las carpetas de chart/embeddings.
        if "/charts/_rels/" in low or "/embeddings/_rels/" in low:
            return True
        return False

    tmp = docx_path + ".tmp"
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for n in names:
            if _is_dropped(n):
                continue
            zout.writestr(n, contents[n])
    os.replace(tmp, docx_path)

    return {
        "dropped_parts": len([n for n in drop_parts if not _is_dropped(n)
                              or True]),
        "dropped_rels": dropped_rels,
        "dropped_overrides": dropped_overrides,
    }


def main():
    if len(sys.argv) > 1:
        src = sys.argv[1]
    else:
        src = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           "reporte_original.docx")
    if not os.path.isfile(src):
        print("ERROR: no se encontro el reporte original: %s" % src)
        print("Uso: python prepare_template.py \"ruta/al/reporte_original.docx\"")
        sys.exit(1)

    doc = Document(src)
    paragraphs = list(doc.paragraphs)
    n = len(paragraphs)
    print("Reporte original con %d parrafos y %d tablas."
          % (n, len(doc.tables)))

    # 1) Capturar referencias antes de mutar.
    def get(i):
        if 0 <= i < n:
            return paragraphs[i]
        return None

    text_targets = {i: get(i) for i in TEXT_PARAGRAPHS if get(i) is not None}
    figure_targets = {i: get(i) for i in FIGURE_PARAGRAPHS
                      if get(i) is not None}
    cons_targets = {name: [get(i) for i in idxs if get(i) is not None]
                    for name, idxs in CONSIDERATION_GROUPS.items()}

    # 2) Texto.
    for idx, tag in TEXT_PARAGRAPHS.items():
        if idx in text_targets:
            set_paragraph_text(text_targets[idx], tag)

    # 3) Figuras.
    for idx, fig in FIGURE_PARAGRAPHS.items():
        if idx in figure_targets:
            set_paragraph_text(figure_targets[idx], "{{ %s }}" % fig)

    # 4) Considerations como bucle Jinja.
    for name, paras in cons_targets.items():
        if not paras:
            continue
        first = paras[0]
        first.insert_paragraph_before("{%%p for c in %s %%}" % name,
                                      style=first.style)
        set_paragraph_text(first, "{{ c }}")
        for middle in paras[1:-1]:
            delete_paragraph(middle)
        set_paragraph_text(paras[-1], "{%p endfor %}")

    # 5) Tablas: 0 consolidada, 1 deliveries, 2 service trucks, 3 weekly,
    #    4 findings (no tocar).
    tables = doc.tables
    if len(tables) >= 1:
        tag_consolidated_table(tables[0])
    if len(tables) >= 2:
        tag_delivery_table(tables[1])
    if len(tables) >= 3:
        tag_trucks_table(tables[2])
    if len(tables) >= 4:
        tag_weekly_table(tables[3])
    # tables[4] -> Findings: queda intacta.

    # 6) Unificar fechas de periodo en captions: "for April, 2026."
    #    "Fuel Outflow Distribution - April 2026.", etc.
    retagged = 0
    for p in doc.paragraphs:
        if retag_period(p):
            retagged += 1
        elif retag_month_year(p):
            retagged += 1
    print("Fechas de periodo unificadas en %d leyendas." % retagged)

    # 7) Fechas en cuadros de texto (portada).
    cover = retag_textbox_dates(doc)
    print("Fechas de portada/cuadros de texto etiquetadas: %d." % cover)

    # 8) Eliminar graficos NATIVOS de Word del cuerpo (la causa del error
    #    'unreadable content'). Todas las figuras del reporte se generan
    #    como imagenes PNG con matplotlib.
    n_charts = strip_native_charts(doc)
    print("Graficos nativos eliminados del cuerpo: %d." % n_charts)

    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       OUTPUT_NAME)
    doc.save(out)

    # 9) Purgar del paquete las partes de chart/embeddings huerfanas + sus
    #    relaciones y Content-Types (cirugia sobre el zip).
    purge = purge_chart_parts(out)
    print("Purga de partes de chart: %s" % purge)

    print("Plantilla generada correctamente: %s" % out)


if __name__ == "__main__":
    main()
